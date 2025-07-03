import pytest
from pathlib import Path
import os
import tempfile
from datetime import datetime, UTC
from src.document_processor import DocumentProcessor

@pytest.fixture
def document_processor():
    return DocumentProcessor()

@pytest.fixture
def sample_text_file():
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write("""Skills:
- Python
- Java
- SQL

Experience:
5 years of software development

Education:
Bachelor's in Computer Science

Certifications:
AWS Certified Developer

Summary:
Experienced software developer with expertise in Python and Java.
""")
    yield f.name
    os.unlink(f.name)

@pytest.fixture
def sample_docx_file():
    # Create a simple docx file for testing
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc = Document()
        doc.add_paragraph("Skills:\n- Python\n- Java\n- SQL")
        doc.add_paragraph("Experience:\n5 years of software development")
        doc.add_paragraph("Education:\nBachelor's in Computer Science")
        doc.add_paragraph("Certifications:\nAWS Certified Developer")
        doc.add_paragraph("Summary:\nExperienced software developer.")
        doc.save(f.name)
    yield f.name
    os.unlink(f.name)

@pytest.fixture
def sample_pdf_file():
    # Create a simple PDF file for testing
    from reportlab.pdfgen import canvas
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
        c = canvas.Canvas(f.name)
        c.drawString(100, 750, "Skills:")
        c.drawString(100, 730, "- Python")
        c.drawString(100, 710, "- Java")
        c.drawString(100, 690, "- SQL")
        c.drawString(100, 650, "Experience:")
        c.drawString(100, 630, "5 years of software development")
        c.save()
    yield f.name
    os.unlink(f.name)

def test_process_directory(document_processor, sample_text_file):
    """Test processing a directory of text files."""
    directory = os.path.dirname(sample_text_file)
    results = document_processor.process_directory(directory)
    assert len(results) > 0
    assert os.path.basename(sample_text_file) in results
    assert "Python" in results[os.path.basename(sample_text_file)]

def test_process_txt(document_processor, sample_text_file):
    """Test processing a text file."""
    content = document_processor._process_txt(sample_text_file)
    assert content is not None
    assert "Python" in content
    assert "Java" in content
    assert "SQL" in content

def test_process_docx(document_processor, sample_docx_file):
    """Test processing a Word document."""
    content = document_processor._process_docx(sample_docx_file)
    assert content is not None
    assert "Python" in content
    assert "Java" in content
    assert "SQL" in content

def test_process_pdf(document_processor, sample_pdf_file):
    """Test processing a PDF file."""
    content = document_processor._process_pdf(sample_pdf_file)
    assert content is not None
    assert "Python" in content
    assert "Java" in content
    assert "SQL" in content

def test_process_txt_with_latin1_encoding(document_processor):
    """Test processing a text file with Latin-1 encoding."""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', encoding='latin-1', delete=False) as f:
        f.write("Test with Latin-1 characters: é, ñ, ü")
    try:
        content = document_processor._process_txt(f.name)
        assert content is not None
        assert "é" in content
        assert "ñ" in content
        assert "ü" in content
    finally:
        os.unlink(f.name)

def test_process_txt_with_encoding_error(document_processor):
    """Test processing a text file with encoding errors."""
    with tempfile.NamedTemporaryFile(mode='wb', suffix='.txt', delete=False) as f:
        f.write(b"Invalid UTF-8: \xff\xfe")
    try:
        content = document_processor._process_txt(f.name)
        # Should fall back to latin-1 and return a string
        assert isinstance(content, str)
        assert "Invalid UTF-8" in content
    finally:
        os.unlink(f.name)

def test_process_docx_with_error(document_processor):
    """Test processing a corrupted DOCX file."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        f.write(b"Not a valid DOCX file")
    try:
        content = document_processor._process_docx(f.name)
        assert content is None
    finally:
        os.unlink(f.name)

def test_process_pdf_with_error(document_processor):
    """Test processing a corrupted PDF file."""
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
        f.write(b"Not a valid PDF file")
    try:
        content = document_processor._process_pdf(f.name)
        assert content is None
    finally:
        os.unlink(f.name)

def test_extract_sections(document_processor):
    """Test extracting sections from document text."""
    text = """Skills:
- Python
- Java
- SQL

Experience:
5 years of software development

Education:
Bachelor's in Computer Science

Certifications:
AWS Certified Developer

Summary:
Experienced software developer.
"""
    sections = document_processor.extract_sections(text)
    assert 'skills' in sections
    assert 'experience' in sections
    assert 'education' in sections
    assert 'certifications' in sections
    assert 'summary' in sections
    assert "Python" in sections['skills']
    assert "5 years" in sections['experience']
    assert "Bachelor's" in sections['education']
    assert "AWS" in sections['certifications']
    assert "Experienced" in sections['summary']

def test_clean_text(document_processor):
    """Test cleaning and normalizing text."""
    text = """  This is a test   text.
    It has multiple    spaces.
    
    And multiple
    newlines.
    
    It also has special characters: @#$%^&*()_+"""
    cleaned = document_processor.clean_text(text)
    assert "  " not in cleaned  # No double spaces
    assert "\n\n\n" not in cleaned  # No triple newlines
    assert "@#$%^&*()_+" not in cleaned  # Special characters removed
    assert "This is a test text" in cleaned

def test_process_document(document_processor, sample_text_file):
    """Test processing a complete document."""
    result = document_processor.process_document(sample_text_file)
    assert result is not None
    assert 'content' in result
    assert 'sections' in result
    assert 'metadata' in result
    assert "Python" in result['content']
    assert result['metadata']['file_type'] == '.txt'
    assert 'filename' in result['metadata']
    assert 'file_size' in result['metadata']
    assert 'created_at' in result['metadata']
    assert 'modified_at' in result['metadata']

def test_process_document_invalid_file(document_processor):
    """Test processing an invalid file."""
    result = document_processor.process_document("nonexistent.txt")
    assert result is None

def test_process_document_unsupported_format(document_processor):
    """Test processing an unsupported file format."""
    with tempfile.NamedTemporaryFile(suffix='.xyz', delete=False) as f:
        pass
    try:
        result = document_processor.process_document(f.name)
        assert result is None
    finally:
        os.unlink(f.name)

def test_extract_metadata(document_processor, sample_text_file):
    """Test extracting file metadata."""
    metadata = document_processor._extract_metadata(Path(sample_text_file))
    assert metadata['filename'] == os.path.basename(sample_text_file)
    assert metadata['file_type'] == '.txt'
    assert isinstance(metadata['file_size'], int)
    assert isinstance(metadata['created_at'], str)
    assert isinstance(metadata['modified_at'], str)

def test_extract_from_pdf(document_processor, sample_pdf_file):
    """Test extracting text from PDF file."""
    text = document_processor._extract_from_pdf(Path(sample_pdf_file))
    assert text is not None
    assert "Python" in text
    assert "Java" in text
    assert "SQL" in text

def test_extract_from_docx(document_processor, sample_docx_file):
    """Test extracting text from DOCX file."""
    text = document_processor._extract_from_docx(Path(sample_docx_file))
    assert text is not None
    assert "Python" in text
    assert "Java" in text
    assert "SQL" in text

def test_extract_from_txt(document_processor, sample_text_file):
    """Test extracting text from TXT file."""
    text = document_processor._extract_from_txt(Path(sample_text_file))
    assert text is not None
    assert "Python" in text
    assert "Java" in text
    assert "SQL" in text

def test_extract_from_pdf_with_error(document_processor):
    """Test extracting text from invalid PDF file."""
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
        f.write(b"Not a valid PDF file")
    try:
        with pytest.raises(Exception):
            document_processor._extract_from_pdf(Path(f.name))
    finally:
        os.unlink(f.name)

def test_extract_from_docx_with_error(document_processor):
    """Test extracting text from invalid DOCX file."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        f.write(b"Not a valid DOCX file")
    try:
        with pytest.raises(Exception):
            document_processor._extract_from_docx(Path(f.name))
    finally:
        os.unlink(f.name)

def test_extract_from_txt_with_error(document_processor):
    """Test extracting text from invalid TXT file."""
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as f:
        f.write(b"Invalid UTF-8: \xff\xfe")
    try:
        with pytest.raises(Exception):
            document_processor._extract_from_txt(Path(f.name))
    finally:
        os.unlink(f.name) 