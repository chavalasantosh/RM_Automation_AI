import os
import pytest
from pathlib import Path
import json
import yaml
import pandas as pd
from docx import Document
from PyPDF2 import PdfWriter
import tempfile
import shutil
from src.enhanced_document_processor import EnhancedDocumentProcessor

@pytest.fixture
def processor():
    """Create a document processor instance for testing."""
    return EnhancedDocumentProcessor()

@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    temp_dir = tempfile.mkdtemp()
    yield temp_dir
    shutil.rmtree(temp_dir)

def create_test_txt(temp_dir: str) -> str:
    """Create a test text file."""
    file_path = os.path.join(temp_dir, "test.txt")
    content = """Summary
Experienced software engineer with expertise in Python and machine learning.

Skills
- Python
- Machine Learning
- Data Analysis

Experience
Senior Software Engineer at Tech Corp (2018-2022)
- Led development of AI-powered applications
- Mentored junior developers

Education
BS in Computer Science, University of Technology (2018)

Certifications
- AWS Certified Developer
- Google Cloud Professional

Projects
- AI Resume Matcher
- Data Analysis Pipeline
"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    return file_path

def create_test_docx(temp_dir: str) -> str:
    """Create a test Word document."""
    file_path = os.path.join(temp_dir, "test.docx")
    doc = Document()
    
    # Add sections
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Experienced software engineer with expertise in Python and machine learning.")
    
    doc.add_heading("Skills", level=1)
    skills = doc.add_paragraph()
    skills.add_run("• Python\n").bold = True
    skills.add_run("• Machine Learning\n").bold = True
    skills.add_run("• Data Analysis").bold = True
    
    doc.add_heading("Experience", level=1)
    exp = doc.add_paragraph()
    exp.add_run("Senior Software Engineer at Tech Corp (2018-2022)\n").bold = True
    exp.add_run("• Led development of AI-powered applications\n")
    exp.add_run("• Mentored junior developers")
    
    doc.save(file_path)
    return file_path

def create_test_pdf(temp_dir: str) -> str:
    """Create a test PDF file."""
    file_path = os.path.join(temp_dir, "test.pdf")
    writer = PdfWriter()
    
    # Add a page with text
    page = writer.add_page()
    page.merge_page(create_test_docx(temp_dir))  # Convert docx to PDF
    
    with open(file_path, "wb") as f:
        writer.write(f)
    return file_path

def create_test_excel(temp_dir: str) -> str:
    """Create a test Excel file."""
    file_path = os.path.join(temp_dir, "test.xlsx")
    
    # Create DataFrame
    data = {
        "Category": ["Summary", "Skills", "Experience", "Education", "Certifications"],
        "Content": [
            "Experienced software engineer...",
            "Python, Machine Learning, Data Analysis",
            "Senior Software Engineer at Tech Corp (2018-2022)",
            "BS in Computer Science, University of Technology (2018)",
            "AWS Certified Developer, Google Cloud Professional"
        ]
    }
    df = pd.DataFrame(data)
    
    # Save to Excel
    df.to_excel(file_path, index=False)
    return file_path

def create_test_csv(temp_dir: str) -> str:
    """Create a test CSV file."""
    file_path = os.path.join(temp_dir, "test.csv")
    
    # Create DataFrame
    data = {
        "Category": ["Summary", "Skills", "Experience", "Education", "Certifications"],
        "Content": [
            "Experienced software engineer...",
            "Python, Machine Learning, Data Analysis",
            "Senior Software Engineer at Tech Corp (2018-2022)",
            "BS in Computer Science, University of Technology (2018)",
            "AWS Certified Developer, Google Cloud Professional"
        ]
    }
    df = pd.DataFrame(data)
    
    # Save to CSV
    df.to_csv(file_path, index=False)
    return file_path

def create_test_html(temp_dir: str) -> str:
    """Create a test HTML file."""
    file_path = os.path.join(temp_dir, "test.html")
    content = """<!DOCTYPE html>
<html>
<head>
    <title>Resume</title>
</head>
<body>
    <h1>Summary</h1>
    <p>Experienced software engineer with expertise in Python and machine learning.</p>
    
    <h1>Skills</h1>
    <ul>
        <li>Python</li>
        <li>Machine Learning</li>
        <li>Data Analysis</li>
    </ul>
    
    <h1>Experience</h1>
    <h2>Senior Software Engineer at Tech Corp (2018-2022)</h2>
    <ul>
        <li>Led development of AI-powered applications</li>
        <li>Mentored junior developers</li>
    </ul>
    
    <h1>Education</h1>
    <p>BS in Computer Science, University of Technology (2018)</p>
    
    <h1>Certifications</h1>
    <ul>
        <li>AWS Certified Developer</li>
        <li>Google Cloud Professional</li>
    </ul>
</body>
</html>"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    return file_path

def create_test_markdown(temp_dir: str) -> str:
    """Create a test Markdown file."""
    file_path = os.path.join(temp_dir, "test.md")
    content = """# Summary
Experienced software engineer with expertise in Python and machine learning.

# Skills
- Python
- Machine Learning
- Data Analysis

# Experience
## Senior Software Engineer at Tech Corp (2018-2022)
- Led development of AI-powered applications
- Mentored junior developers

# Education
BS in Computer Science, University of Technology (2018)

# Certifications
- AWS Certified Developer
- Google Cloud Professional"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    return file_path

def create_test_json(temp_dir: str) -> str:
    """Create a test JSON file."""
    file_path = os.path.join(temp_dir, "test.json")
    content = {
        "summary": "Experienced software engineer with expertise in Python and machine learning.",
        "skills": ["Python", "Machine Learning", "Data Analysis"],
        "experience": {
            "Senior Software Engineer at Tech Corp": {
                "period": "2018-2022",
                "responsibilities": [
                    "Led development of AI-powered applications",
                    "Mentored junior developers"
                ]
            }
        },
        "education": "BS in Computer Science, University of Technology (2018)",
        "certifications": [
            "AWS Certified Developer",
            "Google Cloud Professional"
        ]
    }
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(content, f, indent=2)
    return file_path

def create_test_yaml(temp_dir: str) -> str:
    """Create a test YAML file."""
    file_path = os.path.join(temp_dir, "test.yaml")
    content = {
        "summary": "Experienced software engineer with expertise in Python and machine learning.",
        "skills": ["Python", "Machine Learning", "Data Analysis"],
        "experience": {
            "Senior Software Engineer at Tech Corp": {
                "period": "2018-2022",
                "responsibilities": [
                    "Led development of AI-powered applications",
                    "Mentored junior developers"
                ]
            }
        },
        "education": "BS in Computer Science, University of Technology (2018)",
        "certifications": [
            "AWS Certified Developer",
            "Google Cloud Professional"
        ]
    }
    with open(file_path, "w", encoding="utf-8") as f:
        yaml.dump(content, f, default_flow_style=False)
    return file_path

@pytest.mark.parametrize("file_creator,file_extension", [
    (create_test_txt, ".txt"),
    (create_test_docx, ".docx"),
    (create_test_pdf, ".pdf"),
    (create_test_excel, ".xlsx"),
    (create_test_csv, ".csv"),
    (create_test_html, ".html"),
    (create_test_markdown, ".md"),
    (create_test_json, ".json"),
    (create_test_yaml, ".yaml")
])
def test_process_document(processor, temp_dir, file_creator, file_extension):
    """Test document processing for different file formats."""
    # Create test file
    file_path = file_creator(temp_dir)
    
    # Process document
    result = processor.process_document(file_path)
    
    # Verify result
    assert result is not None
    assert "content" in result
    assert "sections" in result
    assert "metadata" in result
    assert "ai_analysis" in result
    
    # Verify sections
    sections = result["sections"]
    assert "summary" in sections
    assert "skills" in sections
    assert "experience" in sections
    assert "education" in sections
    assert "certifications" in sections
    
    # Verify metadata
    metadata = result["metadata"]
    assert metadata["filename"] == f"test{file_extension}"
    assert "file_size" in metadata
    assert "created_at" in metadata
    assert "modified_at" in metadata
    assert "file_type" in metadata
    assert "mime_type" in metadata

def test_process_batch(processor, temp_dir):
    """Test batch processing of multiple files."""
    # Create test files
    file_paths = [
        create_test_txt(temp_dir),
        create_test_docx(temp_dir),
        create_test_pdf(temp_dir),
        create_test_excel(temp_dir),
        create_test_csv(temp_dir)
    ]
    
    # Process batch
    results = processor.process_batch(file_paths, batch_size=2)
    
    # Verify results
    assert len(results) == len(file_paths)
    for file_path in file_paths:
        assert file_path in results
        assert results[file_path] is not None

def test_save_processed_document(processor, temp_dir):
    """Test saving processed document in different formats."""
    # Create and process a test file
    file_path = create_test_txt(temp_dir)
    content = processor.process_document(file_path)
    
    # Test saving in different formats
    formats = ["json", "yaml", "txt"]
    for format in formats:
        output_path = os.path.join(temp_dir, f"output.{format}")
        assert processor.save_processed_document(content, output_path, format)
        assert os.path.exists(output_path)

def test_invalid_file(processor, temp_dir):
    """Test handling of invalid files."""
    # Test non-existent file
    with pytest.raises(FileNotFoundError):
        processor.process_document(os.path.join(temp_dir, "nonexistent.txt"))
    
    # Test unsupported format
    invalid_file = os.path.join(temp_dir, "test.invalid")
    with open(invalid_file, "w") as f:
        f.write("test")
    with pytest.raises(ValueError):
        processor.process_document(invalid_file)
    
    # Test file too large
    large_file = os.path.join(temp_dir, "large.txt")
    with open(large_file, "w") as f:
        f.write("x" * (processor.max_file_size + 1))
    with pytest.raises(ValueError):
        processor.process_document(large_file)

def test_encoding_handling(processor, temp_dir):
    """Test handling of different file encodings."""
    # Create test files with different encodings
    encodings = ["utf-8", "latin-1", "cp1252", "utf-16"]
    for encoding in encodings:
        file_path = os.path.join(temp_dir, f"test_{encoding}.txt")
        content = "Test content with special characters: é, ñ, 漢字"
        with open(file_path, "w", encoding=encoding) as f:
            f.write(content)
        
        # Process file
        result = processor.process_document(file_path)
        assert result is not None
        assert content in result["content"]

def test_section_extraction(processor, temp_dir):
    """Test extraction of sections from document content."""
    # Create test file with various section formats
    file_path = os.path.join(temp_dir, "sections.txt")
    content = """PROFILE
Experienced software engineer.

TECHNICAL SKILLS
- Python
- Java

WORK EXPERIENCE
Senior Developer at ABC Corp

EDUCATION
BS in Computer Science

CERTIFICATIONS & TRAINING
AWS Certified

PROJECTS & ACHIEVEMENTS
- Project A
- Project B

OTHER INFORMATION
Additional details here."""
    
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    
    # Process document
    result = processor.process_document(file_path)
    sections = result["sections"]
    
    # Verify section extraction
    assert "profile" in sections["summary"].lower()
    assert "technical skills" in sections["skills"].lower()
    assert "work experience" in sections["experience"].lower()
    assert "education" in sections["education"].lower()
    assert "certifications" in sections["certifications"].lower()
    assert "projects" in sections["projects"].lower()
    assert "other information" in sections["other"].lower() 