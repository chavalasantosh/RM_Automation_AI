from typing import List, Dict, Any, Optional
import os
from pathlib import Path
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import logging
from datetime import datetime, timezone
import re
from functools import lru_cache
import logging.config
import shutil
import tempfile

# Configure logging
logging_config = {
    'version': 1,
    'formatters': {
        'standard': {
            'format': '%(asctime)s [%(levelname)s] %(name)s: %(message)s'
        },
    },
    'handlers': {
        'file': {
            'class': 'logging.FileHandler',
            'filename': 'document_processor.log',
            'formatter': 'standard'
        },
        'console': {
            'class': 'logging.StreamHandler',
            'formatter': 'standard'
        }
    },
    'loggers': {
        '': {
            'handlers': ['file', 'console'],
            'level': 'INFO'
        }
    }
}

logging.config.dictConfig(logging_config)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Class for processing different types of documents."""
    
    def __init__(self, max_file_size: int = 10 * 1024 * 1024):  # 10MB default
        """Initialize the document processor."""
        self.logger = logging.getLogger(__name__)
        self.supported_formats = {
            '.txt': self._process_txt,
            '.docx': self._process_docx,
            '.pdf': self._process_pdf
        }
        self.max_file_size = max_file_size
        self.encodings = ['utf-8', 'latin-1', 'cp1252']
        
    def _validate_file(self, file_path: Path) -> bool:
        """Validate file before processing."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if file_path.stat().st_size > self.max_file_size:
            raise ValueError(f"File too large: {file_path}")
            
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
            
        return True
        
    def process_directory(self, directory: str) -> Dict[str, str]:
        """Process all documents in a directory."""
        results = {}
        try:
            for filename in os.listdir(directory):
                filepath = os.path.join(directory, filename)
                suffix = Path(filename).suffix.lower()
                if suffix in self.supported_formats:
                    try:
                        doc_result = self.process_document(filepath)
                        if doc_result and isinstance(doc_result, dict) and 'content' in doc_result:
                            results[filename] = doc_result['content']
                    except Exception as e:
                        self.logger.error(f"Error processing file {filepath}: {str(e)}")
            return results
        except Exception as e:
            self.logger.error(f"Error processing directory {directory}: {str(e)}")
            raise
            
    def _process_txt(self, file_path: str) -> Optional[str]:
        """Process a text file with multiple encoding attempts."""
        for encoding in self.encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        self.logger.error(f"Failed to decode file {file_path} with any supported encoding")
        return None
        
    def _create_temp_copy(self, file_path: Path) -> Path:
        """Create a temporary copy of the file for processing."""
        try:
            temp_dir = Path(tempfile.gettempdir())
            temp_file = temp_dir / f"rme_temp_{file_path.name}"
            shutil.copy2(file_path, temp_file)
            return temp_file
        except Exception as e:
            self.logger.error(f"Error creating temporary copy: {str(e)}")
            raise
            
    def _process_docx(self, file_path: str) -> Optional[str]:
        """Process a Word document with improved error handling."""
        temp_file = None
        try:
            # Create a temporary copy to avoid file locking issues
            temp_file = self._create_temp_copy(Path(file_path))
            
            # Try to open the document
            doc = DocxDocument(temp_file)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            # Also try to extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
                            
            return '\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Error reading Word document {file_path}: {str(e)}")
            return None
            
        finally:
            # Clean up temporary file
            if temp_file and temp_file.exists():
                try:
                    temp_file.unlink()
                except Exception as e:
                    self.logger.warning(f"Error cleaning up temporary file: {str(e)}")
                    
    def _process_pdf(self, file_path: str) -> Optional[str]:
        """Process a PDF file with improved error handling."""
        temp_file = None
        try:
            # Create a temporary copy to avoid file locking issues
            temp_file = self._create_temp_copy(Path(file_path))
            
            # Try to open the PDF
            reader = PdfReader(temp_file)
            
            # Verify PDF is not encrypted
            if reader.is_encrypted:
                self.logger.error(f"PDF file is encrypted: {file_path}")
                return None
                
            # Extract text from each page
            text_parts = []
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    self.logger.warning(f"Error extracting text from page: {str(e)}")
                    continue
                    
            if not text_parts:
                self.logger.warning(f"No text content found in PDF: {file_path}")
                return None
                
            return '\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Error reading PDF file {file_path}: {str(e)}")
            return None
            
        finally:
            # Clean up temporary file
            if temp_file and temp_file.exists():
                try:
                    temp_file.unlink()
                except Exception as e:
                    self.logger.warning(f"Error cleaning up temporary file: {str(e)}")
            
    @lru_cache(maxsize=100)
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a document and return its content and metadata."""
        try:
            file_path = Path(file_path)
            self._validate_file(file_path)
            
            content = self._process_file(file_path)
            if content is None:
                return None
                
            sections = self.extract_sections(content)
            metadata = self._extract_metadata(file_path)
            
            return {
                'content': content,
                'sections': sections,
                'metadata': {
                    **metadata,
                    'processed_at': datetime.now(timezone.utc).isoformat()
                }
            }
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {str(e)}")
            return None
            
    def _process_file(self, file_path: Path) -> Optional[str]:
        """Dispatch to the correct file processor based on file extension."""
        suffix = file_path.suffix.lower()
        processor = self.supported_formats.get(suffix)
        if processor:
            return processor(str(file_path))
        raise ValueError(f"Unsupported file format: {suffix}")
        
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract sections from document text."""
        sections = {
            'skills': '',
            'experience': '',
            'education': '',
            'certifications': '',
            'summary': ''
        }
        
        # Common section headers
        section_patterns = {
            'skills': r'(?i)(?:skills|technical skills|expertise)(?::|$)(.*?)(?=\n\s*\n|\Z)',
            'experience': r'(?i)(?:experience|work experience|employment)(?::|$)(.*?)(?=\n\s*\n|\Z)',
            'education': r'(?i)(?:education|academic|qualification)(?::|$)(.*?)(?=\n\s*\n|\Z)',
            'certifications': r'(?i)(?:certifications|certificates|accreditations)(?::|$)(.*?)(?=\n\s*\n|\Z)',
            'summary': r'(?i)(?:summary|profile|about)(?::|$)(.*?)(?=\n\s*\n|\Z)'
        }
        
        # Extract sections
        for section, pattern in section_patterns.items():
            match = re.search(pattern, text, re.DOTALL)
            if match:
                sections[section] = match.group(1).strip()
                
        return sections
        
    def clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        # Combine all regex operations into a single pass
        patterns = [
            (r'\s+', ' '),
            (r'[^\w\s.,;:()\-/]', ''),
            (r'\n\s*\n', '\n\n')
        ]
        for pattern, replacement in patterns:
            text = re.sub(pattern, replacement, text)
        return text.strip()
        
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from file."""
        stats = file_path.stat()
        return {
            'filename': file_path.name,
            'file_size': stats.st_size,
            'created_at': datetime.fromtimestamp(stats.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stats.st_mtime).isoformat(),
            'file_type': file_path.suffix.lower()
        }
        
    def process_batch(self, file_paths: List[str], batch_size: int = 10) -> Dict[str, Any]:
        """Process multiple files in batches."""
        results = {}
        for i in range(0, len(file_paths), batch_size):
            batch = file_paths[i:i + batch_size]
            for file_path in batch:
                try:
                    result = self.process_document(file_path)
                    if result:
                        results[file_path] = result
                except Exception as e:
                    self.logger.error(f"Error processing {file_path}: {str(e)}")
        return results 