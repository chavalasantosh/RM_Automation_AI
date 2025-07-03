import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import markdown
import html2text

# Create directories if they don't exist
os.makedirs('jd/pdf', exist_ok=True)
os.makedirs('jd/docx', exist_ok=True)
os.makedirs('jd/txt', exist_ok=True)
os.makedirs('jd/html', exist_ok=True)

def save_as_txt(content, filename):
    """Save content as text file."""
    with open(f'jd/txt/{filename}.txt', 'w', encoding='utf-8') as f:
        f.write(content)

def save_as_docx(content, filename):
    """Save content as Word document."""
    doc = Document()
    
    # Add title
    title = doc.add_heading(filename.replace('_', ' ').title(), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Split content into sections and add to document
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            if section.startswith('Job Summary:'):
                doc.add_heading('Job Summary', level=1)
                doc.add_paragraph(section.replace('Job Summary:', '').strip())
            elif section.startswith('Key Responsibilities:'):
                doc.add_heading('Key Responsibilities', level=1)
                for item in section.replace('Key Responsibilities:', '').strip().split('\n'):
                    if item.strip():
                        doc.add_paragraph(item.strip(), style='List Bullet')
            elif section.startswith('Core Skills:'):
                doc.add_heading('Core Skills', level=1)
                for item in section.replace('Core Skills:', '').strip().split('\n'):
                    if item.strip():
                        doc.add_paragraph(item.strip(), style='List Bullet')
            elif section.startswith('Technical Stack:'):
                doc.add_heading('Technical Stack', level=1)
                for item in section.replace('Technical Stack:', '').strip().split('\n'):
                    if item.strip():
                        doc.add_paragraph(item.strip(), style='List Bullet')
            elif section.startswith('Required Skills and Qualifications:'):
                doc.add_heading('Required Skills and Qualifications', level=1)
                for item in section.replace('Required Skills and Qualifications:', '').strip().split('\n'):
                    if item.strip():
                        doc.add_paragraph(item.strip(), style='List Bullet')
            elif section.startswith('Preferred Qualifications:'):
                doc.add_heading('Preferred Qualifications', level=1)
                for item in section.replace('Preferred Qualifications:', '').strip().split('\n'):
                    if item.strip():
                        doc.add_paragraph(item.strip(), style='List Bullet')
            else:
                doc.add_paragraph(section.strip())
    
    doc.save(f'jd/docx/{filename}.docx')

def save_as_pdf(content, filename):
    """Save content as PDF file."""
    pdf = FPDF()
    pdf.add_page()
    
    # Set font
    pdf.set_font("Arial", "B", 16)
    
    # Add title
    pdf.cell(0, 10, filename.replace('_', ' ').title(), ln=True, align='C')
    pdf.ln(10)
    
    # Set font for content
    pdf.set_font("Arial", "", 12)
    
    # Split content into sections and add to PDF
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            if section.startswith('Job Summary:'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, 'Job Summary', ln=True)
                pdf.set_font("Arial", "", 12)
                pdf.multi_cell(0, 10, section.replace('Job Summary:', '').strip())
            elif section.startswith('Key Responsibilities:'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, 'Key Responsibilities', ln=True)
                pdf.set_font("Arial", "", 12)
                for item in section.replace('Key Responsibilities:', '').strip().split('\n'):
                    if item.strip():
                        pdf.multi_cell(0, 10, f"• {item.strip()}")
            elif section.startswith('Core Skills:'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, 'Core Skills', ln=True)
                pdf.set_font("Arial", "", 12)
                for item in section.replace('Core Skills:', '').strip().split('\n'):
                    if item.strip():
                        pdf.multi_cell(0, 10, f"• {item.strip()}")
            elif section.startswith('Technical Stack:'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, 'Technical Stack', ln=True)
                pdf.set_font("Arial", "", 12)
                for item in section.replace('Technical Stack:', '').strip().split('\n'):
                    if item.strip():
                        pdf.multi_cell(0, 10, f"• {item.strip()}")
            elif section.startswith('Required Skills and Qualifications:'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, 'Required Skills and Qualifications', ln=True)
                pdf.set_font("Arial", "", 12)
                for item in section.replace('Required Skills and Qualifications:', '').strip().split('\n'):
                    if item.strip():
                        pdf.multi_cell(0, 10, f"• {item.strip()}")
            elif section.startswith('Preferred Qualifications:'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, 'Preferred Qualifications', ln=True)
                pdf.set_font("Arial", "", 12)
                for item in section.replace('Preferred Qualifications:', '').strip().split('\n'):
                    if item.strip():
                        pdf.multi_cell(0, 10, f"• {item.strip()}")
            else:
                pdf.multi_cell(0, 10, section.strip())
            pdf.ln(5)
    
    pdf.output(f'jd/pdf/{filename}.pdf')

def save_as_html(content, filename):
    """Save content as HTML file."""
    # Convert content to HTML using markdown
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>{filename.replace('_', ' ').title()}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #2c3e50; text-align: center; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        ul {{ list-style-type: none; padding-left: 20px; }}
        li {{ margin: 10px 0; }}
        li:before {{ content: "•"; color: #3498db; font-weight: bold; display: inline-block; width: 1em; margin-left: -1em; }}
    </style>
</head>
<body>
    <h1>{filename.replace('_', ' ').title()}</h1>
"""
    
    # Split content into sections and convert to HTML
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            if section.startswith('Job Summary:'):
                html_content += f"<h2>Job Summary</h2>\n<p>{section.replace('Job Summary:', '').strip()}</p>\n"
            elif section.startswith('Key Responsibilities:'):
                html_content += "<h2>Key Responsibilities</h2>\n<ul>\n"
                for item in section.replace('Key Responsibilities:', '').strip().split('\n'):
                    if item.strip():
                        html_content += f"<li>{item.strip()}</li>\n"
                html_content += "</ul>\n"
            elif section.startswith('Core Skills:'):
                html_content += "<h2>Core Skills</h2>\n<ul>\n"
                for item in section.replace('Core Skills:', '').strip().split('\n'):
                    if item.strip():
                        html_content += f"<li>{item.strip()}</li>\n"
                html_content += "</ul>\n"
            elif section.startswith('Technical Stack:'):
                html_content += "<h2>Technical Stack</h2>\n<ul>\n"
                for item in section.replace('Technical Stack:', '').strip().split('\n'):
                    if item.strip():
                        html_content += f"<li>{item.strip()}</li>\n"
                html_content += "</ul>\n"
            elif section.startswith('Required Skills and Qualifications:'):
                html_content += "<h2>Required Skills and Qualifications</h2>\n<ul>\n"
                for item in section.replace('Required Skills and Qualifications:', '').strip().split('\n'):
                    if item.strip():
                        html_content += f"<li>{item.strip()}</li>\n"
                html_content += "</ul>\n"
            elif section.startswith('Preferred Qualifications:'):
                html_content += "<h2>Preferred Qualifications</h2>\n<ul>\n"
                for item in section.replace('Preferred Qualifications:', '').strip().split('\n'):
                    if item.strip():
                        html_content += f"<li>{item.strip()}</li>\n"
                html_content += "</ul>\n"
            else:
                html_content += f"<p>{section.strip()}</p>\n"
    
    html_content += """</body>
</html>"""
    
    with open(f'jd/html/{filename}.html', 'w', encoding='utf-8') as f:
        f.write(html_content)

def main():
    # Job descriptions
    jobs = {
        'principal_ai_scientist': """Principal AI Scientist (Generative AI & LLMs)
Location: Bangalore, India

Job Summary: We are seeking a visionary Principal AI Scientist to lead our research and development in Generative AI and Large Language Models (LLMs). This is a senior leadership role for an individual with a deep theoretical and practical background in state-of-the-art AI. You will set the strategic direction for our AI initiatives, mentor a team of top-tier researchers and engineers, and drive the creation of transformative AI-powered products that redefine our industry.

Key Responsibilities:
Lead the design, development, and fine-tuning of large-scale language models and generative foundation models.
Architect and implement novel approaches for model alignment, reinforcement learning from human feedback (RLHF), and prompt engineering at scale.
Drive research from conception to publication in top-tier AI conferences (e.g., NeurIPS, ICML, ICLR).
Collaborate with executive leadership to define the AI technology roadmap and its alignment with business goals.
Mentor and guide senior AI/ML engineers and data scientists, fostering a culture of innovation and research excellence.
Oversee the development of scalable MLOps infrastructure for distributed training, inference optimization, and model lifecycle management.
Act as the company's subject matter expert on Generative AI, engaging with the academic community and representing our work at industry events.

Core Skills:
Research Leadership: Guiding research agendas and inspiring breakthrough innovations.
Strategic Thinking: Aligning advanced research with tangible business outcomes.
Mentorship & Coaching: Developing and nurturing top-tier AI talent.
Academic Communication: Publishing and presenting complex research effectively.
Problem Formulation: Deconstructing ambiguous business problems into solvable AI challenges.

Technical Stack:
Languages: Python
Frameworks: PyTorch, TensorFlow/JAX
Models & Architectures: Transformer-based LLMs, Multi-modal Models
Techniques: Reinforcement Learning from Human Feedback (RLHF), Model Quantization, Distillation, Prompt Engineering
MLOps: Kubeflow, MLflow, Weights & Biases
Infrastructure: High-Performance Computing (HPC), GPU/TPU Clusters
Principles: Ethical & Responsible AI

Required Skills and Qualifications:
PhD in Computer Science, AI, or a related field with a focus on deep learning, NLP, or generative models.
10+ years of hands-on experience in AI/ML research and development, with at least 5 years in a leadership or principal role.
Proven track record of publications in premier AI conferences.
Expert-level proficiency in Python and deep learning frameworks such as PyTorch or TensorFlow/JAX.
Demonstrable experience in building, training, and deploying large-scale transformer-based models.
Deep understanding of model quantization, distillation, and optimization techniques for efficient inference.

Preferred Qualifications:
Experience with multi-modal models (text, image, audio).
Familiarity with advanced MLOps tools and platforms (e.g., Kubeflow, MLflow, Weights & Biases).
Experience with high-performance computing (HPC) environments and distributed training on TPUs or large GPU clusters.
Expertise in ethical AI principles and developing responsible AI systems.""",

        'quantum_computing_scientist': """Quantum Computing Scientist
Location: Hyderabad, India

Job Summary: Join our pioneering research group as a Quantum Computing Scientist and contribute to solving some of the world's most complex computational problems. You will be responsible for developing and implementing quantum algorithms, exploring their applications in optimization, simulation, and machine learning. This role is for a highly motivated researcher passionate about pushing the boundaries of computing.

Key Responsibilities:
Design, implement, and test novel quantum algorithms for specific business applications (e.g., financial modeling, drug discovery, materials science).
Work with quantum hardware providers and cloud platforms to execute algorithms on real quantum processors.
Research and develop techniques for quantum error correction and mitigation.
Collaborate with classical computing experts to design hybrid quantum-classical solutions.
Publish research findings in leading scientific journals and present at international conferences.
Contribute to the development of our internal quantum computing software stack and libraries.

Core Skills:
Scientific Research: Applying rigorous scientific methods to computational problems.
Algorithm Design: Inventing and adapting algorithms for quantum architectures.
Theoretical Abstraction: Translating complex physical phenomena into computational models.
Collaborative Research: Working effectively in cross-disciplinary teams.
Complex Problem Solving: Tackling problems that are intractable for classical computers.

Technical Stack:
Languages: Python
Quantum SDKs: Qiskit (IBM), Cirq (Google), PennyLane
Core Concepts: Quantum Mechanics, Quantum Circuits, Quantum Information Theory, Error Correction
Algorithms: Variational Quantum Eigensolver (VQE), Quantum Approximate Optimization Algorithm (QAOA), Quantum Machine Learning (QML)
Hardware: Experience with physical quantum processors.
Techniques: Hybrid Quantum-Classical Methods, Quantum System Simulation

Required Skills and Qualifications:
PhD in Quantum Physics, Quantum Information, Computer Science, or a related field.
Strong theoretical understanding of quantum mechanics, quantum circuits, and quantum information theory.
Proficiency in quantum programming using SDKs like Qiskit (IBM), Cirq (Google), or PennyLane.
Experience in simulating quantum systems and algorithms.
Strong programming skills in Python.
A record of peer-reviewed publications in the field of quantum computing.

Preferred Qualifications:
Experience with specific quantum algorithm families (e.g., VQE, QAOA, Quantum Machine Learning).
Hands-on experience running experiments on physical quantum hardware.
Knowledge of classical optimization and machine learning techniques.
Familiarity with compiler design or low-level hardware control.""",

        'lead_sre': """Lead Site Reliability Engineer (SRE)
Location: Pune, India

Job Summary: As a Lead Site Reliability Engineer, you will be the technical authority for the reliability, scalability, and performance of our most critical distributed systems. This is a leadership role that combines deep software engineering expertise with a passion for operational excellence. You will guide a team of SREs, set technical strategy, and engineer solutions to eliminate toil, automate operations, and ensure our services meet stringent availability and latency goals.

Key Responsibilities:
Architect and implement solutions for ultra-high availability, disaster recovery, and automated scaling across a multi-cloud environment.
Lead the design and implementation of advanced observability platforms using distributed tracing, metrics, and logging at massive scale.
Define Service Level Objectives (SLOs) and error budgets for critical services and drive their adoption across the engineering organization.
Conduct deep-dive post-mortems for high-severity incidents and drive the implementation of preventative measures.
Mentor and develop a team of SREs, fostering a culture of automation, blameless analysis, and continuous improvement.
Engineer custom tooling and automation to eliminate operational toil and improve system efficiency.
Partner with software architects to influence the design of new systems for improved reliability and operability.

Core Skills:
Systems Thinking: Understanding complex systems and their interdependencies.
Technical Leadership: Guiding teams to make sound architectural and operational decisions.
Incident Management: Leading incident response with a calm, methodical approach.
Automation Mindset: A relentless drive to automate manual and repetitive tasks.
Mentorship: Elevating the skills and careers of team members.

Technical Stack:
Languages: Go, Python, Rust
Cloud: AWS, GCP, Azure
Container Orchestration: Kubernetes, Docker
Service Mesh: Istio, Linkerd
Infrastructure as Code (IaC): Terraform, Pulumi
Observability: Prometheus, Grafana, OpenTelemetry
Reliability Techniques: Chaos Engineering (Gremlin, Chaos Toolkit)
Distributed Systems: Raft, Paxos, Large-scale databases (Vitess, CockroachDB, TiDB)

Required Skills and Qualifications:
12+ years of experience in software engineering, with at least 5 years in a senior SRE or DevOps leadership role.
Expert-level knowledge of Kubernetes and its ecosystem (e.g., Service Mesh like Istio or Linkerd, Custom Controllers).
Deep expertise with a major cloud provider (AWS, GCP, or Azure), including their core infrastructure and PaaS offerings.
Proficiency in infrastructure-as-code tools like Terraform or Pulumi.
Strong programming skills in Go, Python, or Rust.
Demonstrable experience in designing and managing large-scale monitoring and observability systems (Prometheus, Grafana, OpenTelemetry).

Preferred Qualifications:
Experience with chaos engineering principles and tools (e.g., Gremlin, Chaos Toolkit).
Deep understanding of Linux kernel internals, networking protocols (TCP/IP, BGP), and distributed systems theory (e.g., consensus algorithms like Raft/Paxos).
Experience managing large-scale database clusters (e.g., Vitess, CockroachDB, TiDB).""",

        'cybersecurity_architect': """Cybersecurity Architect (Zero Trust & Cloud-Native)
Location: Mumbai, India

Job Summary: We are looking for a forward-thinking Cybersecurity Architect to design and implement our next-generation security posture based on Zero Trust principles. In this strategic role, you will be responsible for creating a comprehensive security architecture for our cloud-native and hybrid environments. You will be the technical authority on securing modern infrastructure, applications, and data flows, ensuring security is an integral part of our development and operational lifecycle.

Key Responsibilities:
Design and evangelize a comprehensive Zero Trust security architecture for the entire organization.
Architect and implement security solutions for cloud-native environments, including container security, serverless security, and Kubernetes security.
Develop security-as-code practices, integrating automated security controls and testing into CI/CD pipelines (DevSecOps).
Lead the selection and implementation of advanced security technologies, such as Cloud Security Posture Management (CSPM), Cloud Workload Protection Platforms (CWPP), and Identity-Aware Proxies.
Create threat models for complex distributed systems and design compensating controls.
Define security standards and reference architectures for public cloud deployments (AWS, Azure, GCP).
Act as a senior consultant to engineering teams on secure design patterns and best practices.

Core Skills:
Strategic Security Design: Creating long-term, resilient security roadmaps.
Threat Modeling: Proactively identifying and mitigating potential security flaws.
Risk Assessment: Balancing security controls with business agility and cost.
Communication & Evangelism: Articulating security concepts to technical and non-technical audiences.
Policy Development: Creating clear, actionable security standards and guidelines.

Technical Stack:
Frameworks & Philosophies: Zero Trust, DevSecOps, SASE
Cloud Security: AWS, GCP, Azure, Cloud Security Posture Management (CSPM), Cloud Workload Protection Platforms (CWPP)
Container Security: Kubernetes Security (Pod Security Policies, Network Policies), Service Mesh Security
Identity & Access: OAuth2, OIDC, SAML, SPIFFE/SPIRE, Identity-Aware Proxies
Automation (IaC & Scripting): Terraform, CloudFormation, Python
Advanced Techniques: Confidential Computing, API Security, Security Analytics
Certifications: CISSP-ISSAP, CCSP, GIAC Certified Architect (GCSA)

Required Skills and Qualifications:
10+ years of experience in cybersecurity, with at least 4 years in a security architecture role.
Deep expertise in Zero Trust networking and identity-centric security models.
Strong hands-on experience with securing Kubernetes environments (e.g., Pod Security Policies, Network Policies, Service Mesh security).
In-depth knowledge of cloud security services across AWS, GCP, or Azure.
Experience with Infrastructure as Code (Terraform, CloudFormation) and scripting languages (Python).
Familiarity with modern identity protocols and standards (e.g., OAuth2, OIDC, SAML, SPIFFE/SPIRE).

Preferred Qualifications:
Professional certifications such as CISSP-ISSAP, CCSP, or GIAC Certified Architect (GCSA).
Experience with confidential computing and data encryption in transit, at rest, and in use.
Knowledge of advanced threat detection techniques and security analytics.
Experience in securing API-driven microservices architectures.""",

        'principal_iot_architect': """Principal IoT Solutions Architect (Edge & Digital Twin Focus)
Location: Chennai, India

Job Summary: We are seeking a Principal IoT Solutions Architect to lead the design of our next-generation Industrial IoT (IIoT) platform. This role requires a blend of hardware, software, and cloud expertise to create scalable, end-to-end solutions that leverage Edge AI and Digital Twin technologies. You will be the technical visionary responsible for architecting solutions that connect thousands of devices, process data in real-time, and provide actionable insights to drive operational efficiency and innovation.

Key Responsibilities:
Architect end-to-end IoT solutions, from device connectivity and edge processing to cloud data ingestion and analytics.
Design and implement strategies for Edge Computing and Edge AI, deploying machine learning models on constrained devices.
Lead the development of our Digital Twin framework, creating high-fidelity virtual representations of physical assets and processes.
Select and evaluate IoT hardware, communication protocols (e.g., MQTT, CoAP, LoRaWAN, 5G), and device management platforms.
Ensure the security, scalability, and reliability of the entire IoT ecosystem.
Collaborate with data scientists, product managers, and business stakeholders to translate requirements into technical architecture.
Provide technical leadership and mentorship to a team of IoT engineers.

Core Skills:
End-to-End Architecture: Designing holistic solutions from physical hardware to cloud services.
Solution Design: Translating complex business needs into viable, scalable technical plans.
Technical Vision: Defining the future state of the IoT platform and technology roadmap.
Vendor & Technology Evaluation: Making informed decisions on hardware and software partners.
Stakeholder Management: Communicating effectively across business, product, and engineering teams.

Technical Stack:
Languages: Python, C++, Go
IoT Cloud Platforms: AWS IoT, Azure IoT Hub, Google Cloud IoT Core
Edge Computing: Edge AI, Kubernetes at the Edge (K3s), Docker
Core Concepts: Digital Twin, IoT Security (Device Identity, Secure Boot)
Connectivity Protocols: MQTT, CoAP, LoRaWAN, 5G
Industrial Protocols: Modbus, OPC-UA
Data Stack: Time-series databases (InfluxDB, TimescaleDB), Stream Processing (Kafka, Flink)
Hardware: Embedded Systems, Microcontrollers (MCUs)

Required Skills and Qualifications:
12+ years of experience in software development or systems engineering, with at least 5 years in an IoT-focused architecture role.
Deep expertise with a major IoT cloud platform (AWS IoT, Azure IoT Hub, or Google Cloud IoT Core).
Proven experience in designing and deploying edge computing solutions.
Strong understanding of IoT security principles (device identity, secure boot, data encryption).
Proficiency in at least one programming language like Python, C++, or Go.
Familiarity with containerization on edge devices (Docker, K3s).

Preferred Qualifications:
Experience with Digital Twin platforms or building custom twin solutions.
Knowledge of time-series databases (e.g., InfluxDB, TimescaleDB) and stream processing frameworks (e.g., Kafka, Flink).
Experience with embedded systems and hardware prototyping.
Familiarity with industrial protocols (e.g., Modbus, OPC-UA)."""
    }
    
    # Save each job description in all formats
    for job_id, content in jobs.items():
        print(f"\nProcessing {job_id}...")
        save_as_txt(content, job_id)
        save_as_docx(content, job_id)
        save_as_pdf(content, job_id)
        save_as_html(content, job_id)
        print(f"Saved {job_id} in all formats.")

if __name__ == "__main__":
    main() 